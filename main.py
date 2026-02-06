import tkinter as tk
from tkinter import filedialog, messagebox, Text
from Bio import SeqIO
from io import StringIO
from docx import Document

class FastqToFastaConverter:
    def __init__(self, master):
        self.master = master
        self.master.title("FASTQ to FASTA Converter")
        self.master.geometry("600x500")
        self.master.configure(bg="#1E1E1E")  

        # Configure grid weights for responsiveness
        self.master.grid_rowconfigure(1, weight=1)
        self.master.grid_rowconfigure(2, weight=1)
        self.master.grid_columnconfigure(0, weight=1)

        # Header label
        self.header = tk.Label(self.master, text="FASTQ to FASTA Converter", font=("Helvetica", 16, "bold"), bg="#1E1E1E", fg="white", pady=10)
        self.header.grid(row=0, column=0, sticky="nsew")

        # File content display area
        self.file_content_text = Text(self.master, bg="#3C3C3C", fg="white", font=("Helvetica", 10), wrap="word")
        self.file_content_text.grid(row=1, column=0, sticky="nsew", padx=10, pady=5)

        # Converted FASTA display area
        self.converted_fasta_text = Text(self.master, bg="#3C3C3C", fg="white", font=("Helvetica", 10), wrap="word")
        self.converted_fasta_text.grid(row=2, column=0, sticky="nsew", padx=10, pady=5)

        # Footer frame
        self.footer_frame = tk.Frame(master, bg="#2C2C2C")
        self.footer_frame.grid(row=3, column=0, sticky="nsew", padx=10, pady=5)
        self.footer_frame.grid_columnconfigure((0, 1, 2, 3), weight=1)  # Equal space for buttons

        # Quality score label
        self.quality_score_label = tk.Label(self.footer_frame, text="", font=("Helvetica", 10), bg="#2C2C2C", fg="white")
        self.quality_score_label.grid(row=0, column=0, columnspan=4, pady=5)

        # Select file button
        self.select_file_button = tk.Button(self.footer_frame, text="Select FASTQ File", command=self.select_fastq_file, font=("Helvetica", 12), bg="#4CAF50", fg="white", padx=5, pady=5, borderwidth=0)
        self.select_file_button.grid(row=1, column=0, padx=5, pady=5, sticky="ew")

        # Convert button in footer
        self.convert_button = tk.Button(self.footer_frame, text="Convert to FASTA", command=self.convert_fastq_to_fasta, font=("Helvetica", 12), bg="#4CAF50", fg="white", padx=5, pady=5, borderwidth=0)
        self.convert_button.grid(row=1, column=1, padx=5, pady=5, sticky="ew")

        # Save to Word button
        self.save_button = tk.Button(self.footer_frame, text="Save to Word", command=self.save_to_word, font=("Helvetica", 12), bg="#4CAF50", fg="white", padx=5, pady=5, borderwidth=0)
        self.save_button.grid(row=1, column=2, padx=5, pady=5, sticky="ew")

        # Exit button in footer
        self.exit_button = tk.Button(self.footer_frame, text="Exit", command=master.quit, font=("Helvetica", 12), bg="#f44336", fg="white", padx=5, pady=5, borderwidth=0)
        self.exit_button.grid(row=1, column=3, padx=5, pady=5, sticky="ew")

        self.selected_file = None  # Store the selected file path

    # Other methods remain the same as in your provided code



    def calculate_average_quality(self, quality_scores):
        return sum(quality_scores) / len(quality_scores)

    def quality_filter(self, quality_scores, threshold=20):
        return self.calculate_average_quality(quality_scores) >= threshold

    def select_fastq_file(self):
        self.selected_file = filedialog.askopenfilename(title="Select FASTQ File", filetypes=[("FASTQ files", "*.fastq"), ("All files", "*.*")])
        
        if not self.selected_file:
            messagebox.showwarning("Warning", "No FASTQ file selected!")
            return

        self.display_file_content()

    def display_file_content(self):
        # Clear previous content
        self.file_content_text.delete(1.0, tk.END)
        self.converted_fasta_text.delete(1.0, tk.END)  # Clear the FASTA monitoring area

        try:
            # Read and display the content of the selected FASTQ file
            with open(self.selected_file, "r") as f:
                content = f.read()
                self.file_content_text.insert(tk.END, content)
                self.file_content_text.see(tk.END)  # Scroll to the end of the text widget

            # Calculate and display the quality score
            quality_scores = []
            for record in SeqIO.parse(self.selected_file, "fastq"):
                quality_scores.extend(record.letter_annotations["phred_quality"])

            average_quality = self.calculate_average_quality(quality_scores)
            self.quality_score_label.config(text=f"Average Quality Score: {average_quality:.2f}")
        except Exception as e:
            messagebox.showerror("Error", f"Could not open file: {e}")

    def convert_fastq_to_fasta(self):
        if not self.selected_file:
            messagebox.showwarning("Warning", "No FASTQ file selected!")
            return

        try:
            threshold = 20
            total_quality = 0
            num_records = 0
            filtered_count = 0

            fasta_output = StringIO()  

            # Calculate average quality score
            for record in SeqIO.parse(self.selected_file, "fastq"):
                quality_scores = record.letter_annotations["phred_quality"]
                total_quality += sum(quality_scores)
                num_records += len(quality_scores)

            average_quality = total_quality / num_records if num_records > 0 else 0
            
            # Update quality score label
            self.quality_score_label.config(text=f"Average Quality Score: {average_quality:.2f}")

            # Quality filtering and conversion
            for record in SeqIO.parse(self.selected_file, "fastq"):
                quality_scores = record.letter_annotations["phred_quality"]
                if self.quality_filter(quality_scores, threshold):
                    SeqIO.write(record, fasta_output, "fasta")
                    filtered_count += 1

            # Display the converted FASTA content in the text area
            self.converted_fasta_text.delete(1.0, tk.END)  # Clear previous content
            self.converted_fasta_text.insert(tk.END, fasta_output.getvalue())
            self.converted_fasta_text.see(tk.END)  

            messagebox.showinfo("Success", f"{filtered_count} records converted.")
        except Exception as e:
            messagebox.showerror("Error", f"An error occurred: {e}")

    def save_to_word(self):
        content = self.file_content_text.get(1.0, tk.END)
        quality_info = self.quality_score_label.cget("text")
        fasta_content = self.converted_fasta_text.get(1.0, tk.END)

        if not content.strip() or not fasta_content.strip():
            messagebox.showwarning("Warning", "No content to save!")
            return

        word_file = filedialog.asksaveasfilename(defaultextension=".docx", title="Save to Word Document", filetypes=[("Word files", "*.docx"), ("All files", "*.*")])
        
        if not word_file:
            return
        
        # Create a Word document 
        doc = Document()
        doc.add_heading('FASTQ File Content', level=1)
        doc.add_paragraph(content)
        doc.add_heading('Quality Score', level=2)
        doc.add_paragraph(quality_info)
        doc.add_heading('Converted FASTA Content', level=2)
        doc.add_paragraph(fasta_content)

       
        doc.save(word_file)
        messagebox.showinfo("Success", f"Content saved to '{word_file}'.")


root = tk.Tk()
app = FastqToFastaConverter(root)


root.mainloop()
